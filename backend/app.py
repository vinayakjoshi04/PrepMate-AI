# backend/app.py - PRODUCTION READY (for Render)
from flask import Flask, request, jsonify
from flask_cors import CORS
import json
import re
import traceback
import os
from werkzeug.utils import secure_filename

from huggingfaceService import extract_skills, generate_questions, call_huggingface, extract_json_from_text

app = Flask(__name__)

CORS(app, resources={
    r"/api/*": {
        "origins": [
            "https://prep-mate-ai-eight.vercel.app",
            "https://*.vercel.app",
            "http://localhost:3000",
            "http://localhost:5173",
            "http://localhost:5000",
            "http://127.0.0.1:3000",
            "http://127.0.0.1:5173",
            "http://127.0.0.1:5000",
        ],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "supports_credentials": True,
        "max_age": 3600
    }
})

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE


# ─── Preflight ────────────────────────────────────────────────────────────────

@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", request.headers.get('Origin', '*'))
        response.headers.add("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type, Authorization")
        response.headers.add("Access-Control-Max-Age", "3600")
        return response, 200


# ─── Utilities ────────────────────────────────────────────────────────────────

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def clean_json_response(text):
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    return text.strip()


def repair_json(text):
    """
    Best-effort repair of truncated/malformed JSON.
    1. Strip fences + leading prose.
    2. Direct parse.
    3. Extract outermost balanced { }.
    4. Trim lines from bottom until parseable.
    Returns None if all attempts fail.
    """
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    text = text.strip()

    brace_pos = text.find('{')
    if brace_pos == -1:
        return None
    text = text[brace_pos:]

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    depth = 0
    end_pos = -1
    in_string = False
    escape_next = False
    for i, ch in enumerate(text):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end_pos = i
                break

    if end_pos != -1:
        try:
            return json.loads(text[:end_pos + 1])
        except json.JSONDecodeError:
            pass

    lines = text.splitlines()
    for trim in range(len(lines)):
        attempt = '\n'.join(lines[:len(lines) - trim]).rstrip().rstrip(',')
        open_braces   = attempt.count('{') - attempt.count('}')
        open_brackets = attempt.count('[') - attempt.count(']')
        attempt += ']' * open_brackets + '}' * open_braces
        try:
            return json.loads(attempt)
        except json.JSONDecodeError:
            continue

    return None


def extract_text_from_pdf(file_path):
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return None


def extract_text_from_docx(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return None


def extract_text_from_doc(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return text if text.strip() else None
    except Exception as e:
        try:
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore') or None
        except:
            return None


def extract_resume_text(file_path, filename):
    ext = filename.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext == 'docx':
        return extract_text_from_docx(file_path)
    elif ext == 'doc':
        return extract_text_from_doc(file_path)
    elif ext == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"TXT read error: {e}")
    return None


# ─── Root / Health ────────────────────────────────────────────────────────────

@app.route("/", methods=["GET"])
def root():
    return jsonify({
        "service": "PrepMate-AI Backend",
        "version": "1.0.0",
        "status": "running",
        "endpoints": {
            "health":           "/api/health",
            "create_interview": "/api/create-interview",
            "analyze_answer":   "/api/analyze-answer",
            "batch_analyze":    "/api/batch-analyze-answers",
            "analyze_resume":   "/api/analyze-resume",
            "skill_gap":        "/api/skill-gap"
        }
    }), 200


@app.route("/api/health", methods=["GET"])
def health_check():
    return jsonify({
        "status": "healthy",
        "service": "PrepMate-AI Backend",
        "version": "1.0.0",
        "environment": os.getenv('FLASK_ENV', 'development'),
        "frontend": "https://prep-mate-ai-eight.vercel.app",
        "backend":  "https://prepmate-ai-backend-ckrb.onrender.com"
    }), 200


# ─── Create Interview ─────────────────────────────────────────────────────────
#
#  FIX: Now uses questionsCount, difficulty, focusAreas, and roundName that
#  CreateInterview.js sends in step 3. These were silently ignored before.

@app.route("/api/create-interview", methods=["POST"])
def create_interview():
    try:
        print("\n" + "=" * 60)
        print("📨 CREATE INTERVIEW REQUEST")
        print("=" * 60)

        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400

        required_fields = ["jobTitle", "jobDescription", "experienceLevel", "interviewType"]
        missing_fields  = [f for f in required_fields if not data.get(f)]
        if missing_fields:
            return jsonify({"error": f"Missing required fields: {', '.join(missing_fields)}"}), 400

        job_title        = data["jobTitle"]
        job_description  = data["jobDescription"]
        experience_level = data["experienceLevel"]
        interview_type   = data["interviewType"]

        # Preferences from CreateInterview step 3
        questions_count = int(data.get("questionsCount", 5))
        difficulty      = data.get("difficulty", "mixed")   # easy | mixed | hard
        focus_areas     = data.get("focusAreas", [])        # list of strings
        round_name      = data.get("roundName", "Interview Round")

        print(f"📝 {job_title} | {experience_level} | {interview_type}")
        print(f"🎯 Difficulty: {difficulty} | Questions: {questions_count} | Focus: {focus_areas}")

        # Step 1: Extract skills
        skills_text = extract_skills(job_title, job_description[:2500], experience_level)
        skills      = repair_json(clean_json_response(skills_text))
        if not skills:
            return jsonify({"error": "Failed to extract skills. Please try again."}), 500

        print(f"✅ Skills: {len(skills.get('technicalSkills', []))} tech, "
              f"{len(skills.get('softSkills', []))} soft")

        # Step 2: Generate questions — inject difficulty + focus into context
        difficulty_instruction = {
            "easy":  "Ask foundational, entry-level questions only.",
            "mixed": "Mix of beginner, intermediate, and one advanced question.",
            "hard":  "Ask senior-level questions: complex system design, edge cases, tradeoffs.",
        }.get(difficulty, "Mix of beginner, intermediate, and one advanced question.")

        enriched_skills = json.dumps({
            **skills,
            "_difficulty":      difficulty_instruction,
            "_focusAreas":      ", ".join(focus_areas[:6]) if focus_areas else "",
            "_questionsCount":  questions_count,
            "_roundName":       round_name,
        })

        questions_text = generate_questions(
            job_title, job_description[:2000], experience_level,
            interview_type, enriched_skills
        )
        questions_data = repair_json(clean_json_response(questions_text))
        if not questions_data:
            return jsonify({"error": "Failed to generate questions. Please try again."}), 500

        questions_list = questions_data.get("questions", [])
        # Trim to requested count
        questions_list = questions_list[:questions_count]

        if not questions_list:
            return jsonify({"error": "No questions generated. Please try again."}), 500

        print(f"✅ {len(questions_list)} questions for {round_name}")
        return jsonify({"skills": skills, "questions": questions_list}), 200

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"Server error: {str(e)}"}), 500


# ─── Analyze Single Answer ────────────────────────────────────────────────────
#
#  FIX: Skipped answers return instantly without an AI call.
#  FIX: Returns a graceful fallback (200) instead of 500 so Results.js
#       never crashes mid-analysis — it just shows an estimated score.

@app.route("/api/analyze-answer", methods=["POST"])
def analyze_answer():
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400

        question  = data.get('question', '')
        answer    = data.get('answer', '')
        round_num = data.get('round')
        job_title = data.get('jobTitle', '')
        exp_level = data.get('experienceLevel', '')

        # Return zeroed result for skipped answers without hitting the AI
        if answer.strip() == "[Skipped]":
            return jsonify({
                "score": 0,
                "feedback": ["Question was skipped."],
                "strengths": [],
                "improvements": ["Attempt all questions to receive full AI feedback."],
                "hasExamples": False
            }), 200

        round_name = {1: "Technical Round 1", 2: "Technical Round 2"}.get(round_num, "HR Round")

        prompt = f"""You are an expert interviewer. Score this candidate answer.

Job: {job_title} ({exp_level}) — {round_name}
Question: {question[:300]}
Answer: {answer[:1200]}

Return ONLY valid JSON (no markdown):
{{
  "score": 7,
  "feedback": ["Specific point 1", "Specific point 2", "Specific point 3"],
  "strengths": ["Strength 1", "Strength 2"],
  "improvements": ["Improvement 1", "Improvement 2"],
  "hasExamples": true
}}

Rules: score is 0-10 integer, all arrays 2-3 items, return ONLY JSON."""

        response_text = call_huggingface(prompt, max_tokens=600)
        parsed        = repair_json(response_text)

        if not parsed:
            # Graceful word-count fallback so Results.js keeps going
            wc = len(answer.split())
            return jsonify({
                "score":        min(3 + (wc // 20), 8),
                "feedback":     ["AI scoring temporarily unavailable. Score estimated from answer length."],
                "strengths":    ["Attempted the question"],
                "improvements": ["Add concrete examples and quantify your results"],
                "hasExamples":  False
            }), 200

        return jsonify({
            "score":        max(0, min(10, int(parsed.get('score', 5)))),
            "feedback":     parsed.get('feedback',     []),
            "strengths":    parsed.get('strengths',    []),
            "improvements": parsed.get('improvements', []),
            "hasExamples":  bool(parsed.get('hasExamples', False))
        }), 200

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        # 200 + fallback so Results.js doesn't stop
        return jsonify({
            "score": 5,
            "feedback": ["Server error during analysis. Score is estimated."],
            "strengths": [],
            "improvements": ["Please retry for accurate scoring."],
            "hasExamples": False
        }), 200


# ─── Batch Analyze Answers (NEW) ──────────────────────────────────────────────
#
#  Results.js previously called /analyze-answer once per question sequentially.
#  For 3 rounds × 7 questions = 21 LLM calls chained = guaranteed timeout.
#
#  This endpoint takes ALL answers in one request and scores them in a SINGLE
#  LLM call. One network round-trip, one model call, no timeout chain.
#
#  Results.js now calls this instead of looping over /analyze-answer.

@app.route("/api/batch-analyze-answers", methods=["POST"])
def batch_analyze_answers():
    try:
        print("\n" + "=" * 60)
        print("📨 BATCH ANALYZE ANSWERS")
        print("=" * 60)

        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400

        answers   = data.get('answers', [])
        job_title = data.get('jobTitle', 'the role')
        exp_level = data.get('experienceLevel', 'mid-level')

        if not answers:
            return jsonify({"error": "No answers provided"}), 400

        real_answers = [a for a in answers if (a.get('answer') or '').strip() != '[Skipped]']
        skipped_ids  = {a.get('questionId') for a in answers if (a.get('answer') or '').strip() == '[Skipped]'}

        print(f"📝 {len(real_answers)} real answers | {len(skipped_ids)} skipped")

        # Build compact answer block (300 chars per answer to control tokens)
        answer_lines = []
        for i, a in enumerate(real_answers):
            q = (a.get('question') or '')[:150]
            ans = (a.get('answer')   or '')[:300]
            answer_lines.append(f"[{i+1}] Q: {q}\n    A: {ans}")

        answers_block = "\n\n".join(answer_lines)

        prompt = f"""You are an expert interviewer. Score ALL {len(real_answers)} answers for a {exp_level} {job_title} candidate.

{answers_block}

Return ONLY valid JSON (no markdown):
{{
  "results": [
    {{
      "index": 1,
      "score": 7,
      "feedback": ["Point 1", "Point 2"],
      "strengths": ["Strength 1"],
      "improvements": ["Improvement 1"],
      "hasExamples": true
    }}
  ]
}}

Rules: one entry per answer in order, index starts at 1, score 0-10, return ONLY JSON."""

        # ~130 output tokens per answer
        max_tokens = min(200 + len(real_answers) * 130, 2400)

        print(f"🤖 Batch scoring {len(real_answers)} answers (max_tokens={max_tokens})…")
        response_text = call_huggingface(prompt, max_tokens=max_tokens)
        print(f"📝 Response: {len(response_text)} chars")

        parsed = repair_json(response_text)

        results_map = {}

        if parsed and 'results' in parsed:
            ai_results = parsed['results']
            for i, a in enumerate(real_answers):
                ai = ai_results[i] if i < len(ai_results) else {}
                results_map[a['questionId']] = {
                    "score":        max(0, min(10, int(ai.get('score', 5)))),
                    "feedback":     ai.get('feedback',     ["Answer recorded."]),
                    "strengths":    ai.get('strengths',    ["Attempted the question"]),
                    "improvements": ai.get('improvements', ["Add more specific examples"]),
                    "hasExamples":  bool(ai.get('hasExamples', False)),
                    "skipped":      False,
                }
        else:
            # Word-count fallback for all real answers
            print("⚠️  Batch parse failed — using word-count fallback")
            for a in real_answers:
                wc = len((a.get('answer') or '').split())
                results_map[a['questionId']] = {
                    "score":        min(3 + (wc // 25), 8),
                    "feedback":     ["AI scoring temporarily unavailable. Score estimated."],
                    "strengths":    ["Attempted the question"],
                    "improvements": ["Add concrete examples and quantify results"],
                    "hasExamples":  False,
                    "skipped":      False,
                }

        # Fill skipped answers
        for a in answers:
            if a['questionId'] in skipped_ids:
                results_map[a['questionId']] = {
                    "score":        0,
                    "feedback":     ["Question was skipped."],
                    "strengths":    [],
                    "improvements": ["Attempt all questions for full feedback."],
                    "hasExamples":  False,
                    "skipped":      True,
                }

        # Return in original input order
        ordered = [
            results_map.get(a['questionId'], {
                "score": 5, "feedback": [], "strengths": [],
                "improvements": [], "hasExamples": False, "skipped": False
            })
            for a in answers
        ]

        print(f"✅ Batch complete: {len(ordered)} results")
        return jsonify({"results": ordered}), 200

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"Server error: {str(e)}"}), 500


# ─── Analyze Resume ───────────────────────────────────────────────────────────

@app.route("/api/analyze-resume", methods=["POST"])
def analyze_resume():
    try:
        print("\n" + "=" * 60)
        print("📨 ANALYZE RESUME REQUEST")
        print("=" * 60)

        if 'resume' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files['resume']
        if not file.filename:
            return jsonify({"error": "No file selected"}), 400
        if not allowed_file(file.filename):
            return jsonify({"error": "Invalid file type. Please upload PDF, DOC, or DOCX"}), 400

        filename  = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        resume_text = extract_resume_text(file_path, filename)
        try:
            os.remove(file_path)
        except:
            pass

        if not resume_text or len(resume_text.strip()) < 100:
            return jsonify({"error": "Failed to extract text. Ensure the file is not password-protected."}), 400

        print(f"✅ Extracted {len(resume_text)} chars")

        job_description      = request.form.get('jobDescription', '').strip()
        resume_snippet       = resume_text[:3000]
        resume_snippet_short = resume_text[:2500]

        jd_block = (f"\nJOB DESCRIPTION:\n{job_description[:1500]}\n") if job_description else ""

        # Call 1: Analysis JSON
        analysis_prompt = f"""You are an expert ATS resume reviewer. Analyze this resume.

RESUME:
{resume_snippet}
{jd_block}
Return ONLY this JSON (no markdown):
{{
  "atsScore": 75,
  "sectionFeedback": [
    {{"section": "Contact Information",  "score": 85, "feedback": "Clear contact. Add LinkedIn URL."}},
    {{"section": "Professional Summary", "score": 70, "feedback": "Good but generic. Add keywords."}},
    {{"section": "Work Experience",      "score": 72, "feedback": "Relevant but lacks metrics."}},
    {{"section": "Education",            "score": 80, "feedback": "Clear and well-formatted."}},
    {{"section": "Skills",              "score": 65, "feedback": "List not grouped. Separate technical from soft skills."}}
  ],
  "keywordGaps": ["Agile", "Stakeholder Management", "Data Analysis", "Python", "KPIs"],
  "strengths": ["Clear career progression", "Relevant technical skills", "Good educational background"],
  "improvements": ["Add quantifiable achievements", "Use stronger action verbs", "Expand skills section"]
}}

Rules: atsScore 0-100, sectionFeedback 2-6 items, keywordGaps 5-8, strengths 3-4, improvements 3-5. Return ONLY JSON."""

        print("🤖 Call 1/2: Analysis JSON…")
        r1      = call_huggingface(analysis_prompt, max_tokens=1200)
        analysis = repair_json(r1)
        if not analysis:
            return jsonify({"error": "Failed to parse AI analysis. Please try again."}), 500

        defaults = {"atsScore": 50, "sectionFeedback": [], "keywordGaps": [], "strengths": [], "improvements": []}
        for k in defaults:
            if k not in analysis:
                analysis[k] = defaults[k]

        ats_score = max(0, min(100, int(analysis.get('atsScore', 50))))
        for s in analysis.get('sectionFeedback', []):
            if 'score' in s:
                s['score'] = max(0, min(100, int(s['score'])))

        print(f"✅ Call 1 done | ATS: {ats_score}")

        # Call 2: Improved resume as plain text
        improved_resume = ""
        try:
            improve_prompt = f"""Rewrite this resume in clean ATS-optimised plain text.

RESUME:
{resume_snippet_short}

FIX THESE ISSUES:
{chr(10).join('- ' + i for i in analysis.get('improvements', [])[:3])}

ADD THESE KEYWORDS WHERE APPLICABLE:
{', '.join(analysis.get('keywordGaps', [])[:5])}

Rules: plain text only, standard headings (PROFESSIONAL SUMMARY / WORK EXPERIENCE / EDUCATION / SKILLS), action verbs, no invented info. Return ONLY the resume text."""

            print("🤖 Call 2/2: Improved resume…")
            improved_resume = call_huggingface(improve_prompt, max_tokens=1800)
            improved_resume = re.sub(r'```[a-z]*\n?', '', improved_resume).strip()
            print(f"✅ Call 2 done | {len(improved_resume)} chars")
        except Exception as e:
            print(f"⚠️  Call 2 failed (non-critical): {e}")

        return jsonify({
            "atsScore":        ats_score,
            "sectionFeedback": analysis.get('sectionFeedback', []),
            "keywordGaps":     analysis.get('keywordGaps', []),
            "strengths":       analysis.get('strengths', []),
            "improvements":    analysis.get('improvements', []),
            "improvedResume":  improved_resume,
        }), 200

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"Server error: {str(e)}"}), 500


# ─── Skill Gap ────────────────────────────────────────────────────────────────

@app.route("/api/skill-gap", methods=["POST"])
def skill_gap():
    try:
        print("\n" + "=" * 60)
        print("📨 SKILL GAP REQUEST")
        print("=" * 60)

        job_description = request.form.get('jobDescription', '').strip()
        if not job_description:
            return jsonify({"error": "Job description is required."}), 400

        resume_text = ""
        if 'resume' in request.files:
            file = request.files['resume']
            if file and file.filename:
                if not allowed_file(file.filename):
                    return jsonify({"error": "Invalid file type."}), 400
                filename  = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                resume_text = extract_resume_text(file_path, filename) or ""
                try:
                    os.remove(file_path)
                except:
                    pass

        if not resume_text:
            resume_text = request.form.get('resumeText', '').strip()

        if not resume_text or len(resume_text) < 50:
            return jsonify({"error": "Could not extract resume text. Please try a different file."}), 400

        print(f"✅ JD: {len(job_description)} | Resume: {len(resume_text)} chars")

        prompt = f"""You are an expert career coach. Compare the resume against the job description.

JOB DESCRIPTION:
{job_description[:2500]}

RESUME:
{resume_text[:2500]}

Return ONLY valid JSON (no markdown):
{{
  "presentSkills": ["Skill A", "Skill B"],
  "missingSkills": ["Skill X", "Skill Y"],
  "partialSkills": ["Skill P"],
  "summary": "2-3 sentence summary of fit and key gaps.",
  "totalTimeframe": "10-14 weeks",
  "roadmap": [
    {{
      "skill": "Skill Name",
      "priority": "high",
      "timeframe": "2-3 weeks",
      "matchScore": 80,
      "description": "Why this skill matters and how to learn it.",
      "resources": [{{"name": "Resource", "url": "https://example.com"}}],
      "subtasks": ["Task 1", "Task 2", "Task 3"]
    }}
  ]
}}

Rules: roadmap max 6 items, priority = high/medium/low, return ONLY JSON."""

        print("🤖 Skill gap AI call…")
        response_text = call_huggingface(prompt, max_tokens=2000)
        analysis = repair_json(response_text)
        if not analysis:
            return jsonify({"error": "Failed to parse AI response. Please try again."}), 500

        required_keys = ["presentSkills", "missingSkills", "roadmap"]
        missing_keys  = [k for k in required_keys if k not in analysis]
        if missing_keys:
            return jsonify({"error": f"Incomplete AI response (missing: {', '.join(missing_keys)})."}), 500

        for item in analysis.get("roadmap", []):
            if item.get("priority") not in {"high", "medium", "low"}:
                item["priority"] = "medium"
            if "matchScore" in item:
                item["matchScore"] = max(0, min(100, int(item["matchScore"])))

        result = {
            "presentSkills":  analysis.get("presentSkills", []),
            "missingSkills":  analysis.get("missingSkills", []),
            "partialSkills":  analysis.get("partialSkills", []),
            "summary":        analysis.get("summary", ""),
            "totalTimeframe": analysis.get("totalTimeframe", ""),
            "roadmap":        analysis.get("roadmap", [])
        }
        print(f"✅ Skill gap done | Present: {len(result['presentSkills'])} | "
              f"Missing: {len(result['missingSkills'])} | Roadmap: {len(result['roadmap'])}")
        return jsonify(result), 200

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"Server error: {str(e)}"}), 500


# ─── Error Handlers ───────────────────────────────────────────────────────────

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        "error": "Endpoint not found",
        "available_endpoints": [
            "/api/health", "/api/create-interview",
            "/api/analyze-answer", "/api/batch-analyze-answers",
            "/api/analyze-resume", "/api/skill-gap"
        ]
    }), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({"error": "Internal server error", "message": str(error)}), 500

@app.errorhandler(413)
def file_too_large(error):
    return jsonify({"error": "File too large. Maximum size is 5MB."}), 413


# ─── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = int(os.getenv('PORT', 5000))
    env  = os.getenv('FLASK_ENV', 'development')
    print(f"\n🚀 PrepMate-AI Backend | Port {port} | Env {env}\n")
    app.run(debug=(env == 'development'), port=port, host='0.0.0.0')